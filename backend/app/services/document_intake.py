from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path

from pypdf import PdfReader

from app.domain.models import FileIntakeStatus, PageText


@dataclass
class DocumentIntakeResult:
    content_sha256: str
    intake_status: FileIntakeStatus
    page_count: int = 0
    page_texts: list[PageText] = field(default_factory=list)
    is_scanned: bool = False
    precheck_warning: str | None = None
    validation_error: str | None = None

    @property
    def extracted_text(self) -> str:
        return "\n".join(item.text for item in self.page_texts if item.text).strip()


class DocumentIntakeService:
    _SUPPORTED_SUFFIXES = {
        ".pdf",
        ".docx",
        ".pptx",
        ".txt",
        ".md",
        ".csv",
        ".json",
        ".png",
        ".jpg",
        ".jpeg",
        ".bmp",
        ".gif",
        ".tif",
        ".tiff",
        ".webp",
    }
    _IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff", ".webp"}
    _IRRELEVANT_TERMS = (
        "合同",
        "发票",
        "报价单",
        "采购单",
        "部署文档",
        "安装说明",
        "运维手册",
        "产品说明书",
        "用户手册",
        "license agreement",
        "invoice",
        "deployment guide",
        "product manual",
    )

    def __init__(self, *, max_upload_bytes: int, max_pdf_pages: int) -> None:
        self.max_upload_bytes = max_upload_bytes
        self.max_pdf_pages = max_pdf_pages

    def preflight(self, *, filename: str, content_type: str, content: bytes) -> DocumentIntakeResult:
        digest = hashlib.sha256(content).hexdigest()
        suffix = Path(filename).suffix.lower()
        if not content:
            return self._invalid(digest, "文件内容为空。")
        if len(content) > self.max_upload_bytes:
            return self._invalid(digest, "文件超过允许的大小限制。")
        if suffix not in self._SUPPORTED_SUFFIXES:
            return self._invalid(digest, "暂不支持该文件格式。")

        try:
            if suffix == ".pdf":
                result = self._pdf_result(digest, content)
            elif suffix == ".docx":
                result = self._docx_result(digest, content)
            elif suffix in self._IMAGE_SUFFIXES:
                result = DocumentIntakeResult(
                    content_sha256=digest,
                    intake_status=FileIntakeStatus.uploaded,
                    page_count=1,
                    is_scanned=True,
                )
            elif suffix in {".txt", ".md", ".csv", ".json"} or content_type.startswith("text/"):
                text = content.decode("utf-8-sig", errors="replace")
                result = DocumentIntakeResult(
                    content_sha256=digest,
                    intake_status=FileIntakeStatus.uploaded,
                    page_count=1,
                    page_texts=[PageText(page=1, text=text)],
                )
            else:
                # PPTX remains accepted for compatibility and is analyzed later.
                result = DocumentIntakeResult(
                    content_sha256=digest,
                    intake_status=FileIntakeStatus.uploaded,
                    page_count=0,
                )
        except Exception:
            return self._invalid(digest, "文件损坏或无法读取。")

        self._apply_irrelevant_hint(result, filename)
        return result

    def _pdf_result(self, digest: str, content: bytes) -> DocumentIntakeResult:
        reader = PdfReader(BytesIO(content))
        if len(reader.pages) > self.max_pdf_pages:
            return self._invalid(digest, "PDF 页数超过允许的上限。")
        layout_pages = self._pdf_layout_pages(content, len(reader.pages))
        pages: list[PageText] = []
        readable_pages = 0
        for page_number, page in enumerate(reader.pages, start=1):
            plain_text = (page.extract_text() or "").strip()
            layout_text = layout_pages.get(page_number, "")
            text = (
                layout_text
                if self._prefer_layout_text(plain_text, layout_text)
                else plain_text
            )
            pages.append(PageText(page=page_number, text=text))
            if len(re.sub(r"\s+", "", text)) >= 40:
                readable_pages += 1
        return DocumentIntakeResult(
            content_sha256=digest,
            intake_status=FileIntakeStatus.uploaded,
            page_count=len(reader.pages),
            page_texts=pages,
            is_scanned=bool(reader.pages) and readable_pages == 0,
        )

    @classmethod
    def _pdf_layout_pages(
        cls,
        content: bytes,
        page_count: int,
    ) -> dict[int, str]:
        """Extract coordinate-aware text without making medical judgments.

        PDF content streams frequently store table result values separately from
        their labels.  Plain extraction then moves all values to the end of the
        page.  Coordinate-aware extraction restores the visible row relationship
        before the text is sent to the document model.
        """
        try:
            import pdfplumber

            extracted: dict[int, str] = {}
            with pdfplumber.open(BytesIO(content)) as document:
                for page_number, page in enumerate(
                    document.pages[:page_count],
                    start=1,
                ):
                    text = page.extract_text(
                        layout=True,
                        x_tolerance=2,
                        y_tolerance=8,
                    ) or ""
                    normalized = cls._normalize_layout_text(text)
                    if normalized:
                        extracted[page_number] = normalized
            return extracted
        except Exception:
            # The existing pypdf extraction remains the safe fallback for a PDF
            # whose coordinate data cannot be parsed.
            return {}

    @staticmethod
    def _normalize_layout_text(text: str) -> str:
        lines: list[str] = []
        for raw_line in text.splitlines():
            line = raw_line.strip()
            if not line:
                continue
            # Keep visible column boundaries explicit while avoiding thousands
            # of padding spaces from layout-mode extraction.
            line = re.sub(r"[\t ]{2,}", " | ", line)
            lines.append(line)
        return "\n".join(lines)

    @staticmethod
    def _prefer_layout_text(plain_text: str, layout_text: str) -> bool:
        if not layout_text:
            return False
        plain_length = len(re.sub(r"\s+", "", plain_text))
        layout_length = len(re.sub(r"\s+", "", layout_text))
        if layout_length < 40:
            return False
        if plain_length < 40:
            return True
        # Only replace otherwise readable text when the layout extractor found
        # several genuine multi-column rows.  Narrative pages retain the prior
        # pypdf behavior.
        table_rows = sum(
            1
            for line in layout_text.splitlines()
            if line.count(" | ") >= 2
        )
        return table_rows >= 3 and layout_length >= int(plain_length * 0.7)

    def _docx_result(self, digest: str, content: bytes) -> DocumentIntakeResult:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required for DOCX intake") from exc
        document = Document(BytesIO(content))
        lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                line = " | ".join(cell.text.strip() for cell in row.cells)
                if re.sub(r"[|\s]", "", line):
                    lines.append(line)
        text = "\n".join(lines)
        return DocumentIntakeResult(
            content_sha256=digest,
            intake_status=FileIntakeStatus.uploaded,
            page_count=1,
            page_texts=[PageText(page=1, text=text)],
            is_scanned=False,
        )

    def _apply_irrelevant_hint(self, result: DocumentIntakeResult, filename: str) -> None:
        if result.intake_status == FileIntakeStatus.invalid:
            return
        sample = " ".join([filename, *(page.text[:3000] for page in result.page_texts[:3])]).lower()
        if any(term.lower() in sample for term in self._IRRELEVANT_TERMS):
            result.intake_status = FileIntakeStatus.suspected_irrelevant
            result.precheck_warning = "疑似与病例分析无关，请确认是否误传。"

    def _invalid(self, digest: str, message: str) -> DocumentIntakeResult:
        return DocumentIntakeResult(
            content_sha256=digest,
            intake_status=FileIntakeStatus.invalid,
            validation_error=message,
        )
